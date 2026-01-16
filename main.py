import os
import sys
import argparse
from pathlib import Path
from pypdf import PdfReader
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from docx.shared import Inches
from pdf2docx import Converter
import io

# Configuration
SOURCE_DIR = Path("source_pdf")
OUTPUT_MD_DIR = Path("output_md")
OUTPUT_DOCX_DIR = Path("output_docx")
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\Program Files\poppler-25.07.0\Library\bin"
TESSDATA_LANG = "pol"

# Set pytesseract path
if os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
else:
    print(f"Warning: Tesseract not found at {TESSERACT_PATH}. OCR might fail.")

def get_text_extraction(pdf_path):
    """Try to extract text layer from PDF."""
    try:
        reader = PdfReader(pdf_path)
        full_text = ""
        for page in reader.pages:
            full_text += page.extract_text() or ""
        return full_text.strip()
    except Exception as e:
        print(f"Error reading text layer from {pdf_path}: {e}")
        return ""

def get_ocr_extraction(pdf_path):
    """Perform OCR on PDF pages and return text."""
    try:
        images = convert_from_path(pdf_path, poppler_path=POPPLER_PATH)
        full_text = ""
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image, lang=TESSDATA_LANG)
            full_text += f"## Page {i+1}\n\n{text}\n\n"
        return full_text.strip()
    except Exception as e:
        print(f"Error performing OCR on {pdf_path}: {e}")
        return ""

def save_as_md(pdf_path, text, mode, output_dir):
    output_filename = pdf_path.stem + ".md"
    output_path = output_dir / output_filename
    
    md_content = f"# {pdf_path.name}\n\n"
    md_content += f"*Extracted via: {mode} mode*\n\n"
    md_content += "---\n\n"
    md_content += text
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"  - Success! Saved MD to {output_path}")

def save_as_docx_editable(pdf_path, output_dir):
    """
    Convert PDF to editable DOCX preserving layout, images, and text.
    Uses pdf2docx library.
    """
    try:
        output_filename = pdf_path.stem + ".docx"
        output_path = output_dir / output_filename
        
        print(f"  - Converting to editable DOCX (layout + text)...")
        cv = Converter(str(pdf_path))
        cv.convert(str(output_path))
        cv.close()
        print(f"  - Success! Saved Editable DOCX to {output_path}")
    except Exception as e:
        print(f"  - Error saving editable DOCX for {pdf_path.name}: {e}")

def save_as_docx_image(pdf_path, output_dir):
    """
    High-fidelity DOCX export: converts PDF pages to images.
    Preserves exact visual look but text is NOT editable.
    """
    try:
        output_filename = pdf_path.stem + "_image.docx"
        output_path = output_dir / output_filename
        
        doc = Document()
        doc.add_heading(pdf_path.name, 0)
        
        print(f"  - Converting PDF to images for static DOCX preservation...")
        images = convert_from_path(pdf_path, poppler_path=POPPLER_PATH, dpi=200)
        
        for i, image in enumerate(images):
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            
            doc.add_picture(img_byte_arr, width=Inches(6.0))
            if i < len(images) - 1:
                doc.add_page_break()
                
        doc.save(output_path)
        print(f"  - Success! Saved Image-based DOCX to {output_path}")
    except Exception as e:
        print(f"  - Error saving image-based DOCX for {pdf_path.name}: {e}")

def process_pdf(pdf_path, convert_md, convert_docx, convert_docx_image):
    print(f"Processing: {pdf_path.name}")
    
    # MD Conversion logic
    if convert_md:
        extracted_text = get_text_extraction(pdf_path)
        reader = PdfReader(pdf_path)
        num_pages = len(reader.pages)
        is_ocr_needed = len(extracted_text) < (50 * num_pages)
        
        if is_ocr_needed:
            print(f"  - MD: Low text content. Using OCR mode...")
            final_text = get_ocr_extraction(pdf_path)
            mode = "OCR"
        else:
            print(f"  - MD: Text layer detected. Using standard extraction.")
            final_text = ""
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                final_text += f"## Page {i+1}\n\n{page_text}\n\n"
            mode = "Text"
            
        if final_text:
            save_as_md(pdf_path, final_text, mode, OUTPUT_MD_DIR)
            
    # Editable DOCX Conversion logic
    if convert_docx:
        save_as_docx_editable(pdf_path, OUTPUT_DOCX_DIR)

    # Image-based DOCX Conversion logic
    if convert_docx_image:
        save_as_docx_image(pdf_path, OUTPUT_DOCX_DIR)

def main():
    parser = argparse.ArgumentParser(description="PDF to MD/DOCX Converter")
    parser.add_argument("-md", action="store_true", help="Generate Markdown output")
    parser.add_argument("-docx", action="store_true", help="Generate editable DOCX output (preserves layout/text)")
    parser.add_argument("-docx_image", action="store_true", help="Generate image-based DOCX output (preserves visual layout exactly)")
    
    args = parser.parse_args()
    
    # Default to MD if no arguments provided
    if not (args.md or args.docx or args.docx_image):
        print("No format specified. Defaulting to -md.")
        args.md = True
        
    SOURCE_DIR.mkdir(exist_ok=True)
    OUTPUT_MD_DIR.mkdir(exist_ok=True)
    OUTPUT_DOCX_DIR.mkdir(exist_ok=True)
    
    pdf_files = list(SOURCE_DIR.glob("*.pdf"))
    
    if not pdf_files:
        print(f"No PDF files found in {SOURCE_DIR}")
        return

    formats = []
    if args.md: formats.append("MD")
    if args.docx: formats.append("DOCX (Editable)")
    if args.docx_image: formats.append("DOCX (Image)")

    print(f"Found {len(pdf_files)} PDF files. Target formats: {', '.join(formats)}\n")
    
    for pdf_file in pdf_files:
        process_pdf(pdf_file, args.md, args.docx, args.docx_image)

if __name__ == "__main__":
    main()
